from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "casos-clinicos-didaticos" / "manual_interpretacao_cec_perfuselab.docx"


BLUE = RGBColor(30, 64, 175)
DARK = RGBColor(17, 24, 39)
MUTED = RGBColor(75, 85, 99)
LIGHT_BLUE = "DBEAFE"
LIGHT_AMBER = "FEF3C7"
LIGHT_GREEN = "DCFCE7"
LIGHT_RED = "FEE2E2"
LIGHT_GRAY = "F3F4F6"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(8.5)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers, rows, widths=None, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        shade_cell(header_cells[i], header_fill)
        set_cell_text(header_cells[i], header, bold=True, color=DARK)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_callout(doc, title, body, fill=LIGHT_AMBER):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = DARK
    p.add_run("\n" + body).font.size = Pt(9)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = BLUE if level <= 2 else DARK
    return p


def add_reference(doc, text, url):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text + " ")
    r = p.add_run(url)
    r.font.color.rgb = BLUE
    r.underline = True


def set_document_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10)
    normal.font.color.rgb = DARK

    for style_name, size, bold in [
        ("Title", 24, True),
        ("Heading 1", 16, True),
        ("Heading 2", 13, True),
        ("Heading 3", 11, True),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = BLUE if "Heading" in style_name else DARK

    for style_name in ["List Bullet", "List Number"]:
        styles[style_name].font.name = "Calibri"
        styles[style_name].font.size = Pt(10)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PerfuseLab - Manual didatico de interpretacao de casos em CEC")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def build_manual():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    add_footer(section)
    set_document_styles(doc)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PERFUSELAB")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BLUE

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Manual Didatico de Interpretacao de Casos em CEC")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = DARK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Da fisiologia basica a leitura dos alertas, graficos e casos clinicos simulados")
    r.font.size = Pt(13)
    r.font.color.rgb = MUTED

    doc.add_paragraph()
    add_callout(
        doc,
        "Como usar este manual",
        "Leia primeiro o fluxo de interpretacao. Depois abra um CSV no PerfuseLab, tente explicar a tabela e os graficos sozinho, e so entao confira o gabarito TXT do mesmo caso.",
        LIGHT_BLUE,
    )
    add_callout(
        doc,
        "Aviso importante",
        "Este material e educacional. Os casos sao simulados. O PerfuseLab organiza calculos e alertas, mas nao substitui protocolo institucional, supervisao, julgamento clinico ou decisao da equipe assistencial.",
        LIGHT_RED,
    )
    doc.add_paragraph("Versao: junho de 2026")
    doc.add_page_break()

    add_heading(doc, "Sumario Rapido", 1)
    add_numbered(
        doc,
        [
            "Comece pela qualidade do dado: unidade, tempo, SaO2 medida, Hb/Hct, fluxo ou IC.",
            "Calcule a superficie corporal e pense sempre em valores indexados.",
            "Leia iDO2, lactato, SvO2, O2ER, DO2/VO2, gap PCO2 e acido-base em conjunto.",
            "Compare valor absoluto e tendencia. Um numero isolado quase nunca conta a historia inteira.",
            "Entenda que o risco integrado do PerfuseLab e uma regra local, nao um escore validado.",
            "Use os 20 casos como treino progressivo: primeiro os padroes simples, depois os casos discordantes.",
        ],
    )

    add_heading(doc, "A IA Utilizada", 2)
    doc.add_paragraph(
        "No estado atual do projeto, a analise exibida pelo PerfuseLab no navegador nao chama um modelo de IA externo. "
        "Os calculos e alertas sao determinisiticos, escritos em JavaScript, com formulas e limiares documentados. "
        "Este manual e os casos foram produzidos com apoio do Codex/OpenAI durante o desenvolvimento, mas o comportamento do framework e governado pelo codigo local."
    )

    add_heading(doc, "Mapa Mental de 10 Minutos", 1)
    add_table(
        doc,
        ["Etapa", "Pergunta", "O que observar"],
        [
            ["1. Dado", "Posso confiar na amostra?", "Mesmo tempo de CEC, unidade correta, SaO2 medida ou iDO2 informado, Hb/Hct coerentes."],
            ["2. Oferta", "O iDO2 esta adequado?", "iDO2 final, pior ponto, tempo abaixo de 280 e AUC de deficit."],
            ["3. Demanda", "O consumo/extracao subiu?", "iVO2, O2ER, DO2/VO2 e SvO2."],
            ["4. Metabolismo", "O lactato esta subindo?", "Lactato final, razao final/inicial, pH, HCO3, BE, anion gap."],
            ["5. Fluxo regional", "Ha discordancia?", "Gap PCO2, PAM, SvO2 preservada com lactato/gap alterados."],
            ["6. Contexto", "O que o numero nao sabe?", "Temperatura, tempo, estrategia alfa-stat, hemodiluicao, eletrólitos, anticoagulacao, cirurgia."],
        ],
        widths=[1.3, 1.9, 4.4],
    )

    add_heading(doc, "Primeiro: Qualidade do Dado", 1)
    doc.add_paragraph(
        "Antes de interpretar fisiologia, confirme se a tabela e matematicamente auditavel. Muitos erros parecem clinicos, mas nascem de unidade errada, coleta fora de tempo ou campo ausente."
    )
    add_bullets(
        doc,
        [
            "Tempo: cada linha deve representar um momento real da CEC.",
            "Lactato: o PerfuseLab trabalha internamente em mmol/L; quando vier em mg/dL, divide por 9,009.",
            "SaO2: se estiver medida, o sistema recalcula iDO2. Se faltar, o iDO2 pode ser aceito como informado, mas perde auditoria independente.",
            "Fluxo e IC: o fluxo total precisa ser indexado pela BSA. Em BSA alta, fluxo em L/min pode enganar.",
            "Hb/Hct: se apenas um campo existir, o sistema pode estimar o outro por regra simples, mas isso nao substitui laboratorio.",
            "PaO2: contribui pouco para CaO2 em comparacao com Hb e SaO2, mas ajuda na auditoria de oxigenacao e FiO2.",
        ],
    )
    add_callout(
        doc,
        "Regra de ouro",
        "Nunca corrija uma historia clinica sem antes corrigir a planilha. Se SaO2, Hb, fluxo e PaO2 nao pertencem ao mesmo tempo, o iDO2 calculado perde sentido.",
        LIGHT_GREEN,
    )

    add_heading(doc, "Formulas Fundamentais", 1)
    add_table(
        doc,
        ["Metrica", "Formula usada no estudo", "Como interpretar"],
        [
            ["BSA", "sqrt(peso x altura / 3600)", "Base para indexar fluxo e oferta. Mosteller."],
            ["IC", "fluxo total / BSA", "Fluxo em L/min/m2. Ajuda a comparar pacientes de tamanhos diferentes."],
            ["CaO2", "Hb x 1,36 x SaO2 + PaO2 x 0,003", "Conteudo arterial de oxigenio em mL/dL. SaO2 entra como fracao."],
            ["iDO2", "10 x IC x CaO2", "Oferta de oxigenio indexada, em mL/min/m2."],
            ["O2ER", "iVO2 / iDO2 x 100", "Percentual extraido. Sobe quando a demanda consome mais da oferta."],
            ["DO2/VO2", "iDO2 / iVO2", "Inverso pratico da extracao. Abaixo de 3,33 equivale a O2ER >= 30%."],
            ["Lactato", "mg/dL / 9,009 = mmol/L", "Use tendencia e contexto, nao apenas ponto final."],
            ["AUC deficit", "Area entre alvo e iDO2 quando iDO2 < alvo", "Resume profundidade e duracao do deficit."],
        ],
        widths=[1.3, 2.4, 3.9],
    )

    add_heading(doc, "Como Ler a Oferta de Oxigenio", 1)
    doc.add_paragraph(
        "O iDO2 e a peca central da goal-directed perfusion. No PerfuseLab, o alvo fixo de referencia e 280 mL/min/m2, baseado no ensaio de Ranucci. Esse alvo e operacional, nao uma fronteira biologica absoluta."
    )
    add_table(
        doc,
        ["Faixa no PerfuseLab", "Leitura pratica", "Cuidado"],
        [
            [">= 280 mL/min/m2", "Adequado para o alvo GDP fixo.", "Ainda avaliar lactato, gap PCO2, PAM e contexto."],
            ["260 a 279 mL/min/m2", "Limitrofe para GDP.", "Olhar tendencia, tempo abaixo e reserva metabolica."],
            ["< 260 mL/min/m2", "Abaixo da faixa desejada no cartao.", "Se tambem < 90% de 280, gera alerta alto na analise."],
            ["iDO2 informado", "Aceito quando falta SaO2.", "Nao e recalculavel; confirmar origem e alinhamento temporal."],
        ],
        widths=[1.8, 2.6, 3.2],
    )

    add_heading(doc, "Extracao, Consumo e SvO2", 1)
    doc.add_paragraph(
        "A oferta isolada nao diz quanto o organismo esta consumindo. Por isso, o PerfuseLab tambem recalcula O2ER e DO2/VO2 quando iVO2 esta disponivel."
    )
    add_table(
        doc,
        ["Marcador", "Faixa de alerta local", "Mensagem didatica"],
        [
            ["O2ER", ">= 30% moderado; > 40% alto", "Extracao crescente pode indicar reserva menor ou demanda maior."],
            ["DO2/VO2", "<= 3,33 moderado; < 2,5 alto", "E a mesma ideia da extracao, vista pela relacao oferta/consumo."],
            ["SvO2", "<= 70% moderado; < 65% alto", "Sinal global util, mas pode mascarar distribuicao regional."],
            ["iVO2", "Sem corte isolado no app", "Interprete junto com temperatura, anestesia, fluxo e extracao."],
        ],
        widths=[1.5, 2.2, 3.9],
    )
    add_callout(
        doc,
        "SvO2 normal nao encerra o caso",
        "O framework mostra alerta informativo quando SvO2 esta preservada, mas lactato, gap PCO2 ou extracao sugerem discordancia. Isso nao e diagnostico de hipoxia oculta; e um lembrete para olhar a tabela inteira.",
        LIGHT_AMBER,
    )

    add_heading(doc, "Lactato, Gap PCO2 e Acido-Base", 1)
    doc.add_paragraph(
        "Lactato e um marcador poderoso, mas pouco especifico. Durante CEC, ele pode subir por hipoperfusao, resposta adrenergica, hemodiluicao, hiperglicemia, inflamacao, temperatura, fluxo regional ou menor depuracao."
    )
    add_table(
        doc,
        ["Achado", "Como o PerfuseLab alerta", "Como estudar"],
        [
            ["Lactato final 2 a 4 mmol/L", "Moderado", "Comparar com iDO2, O2ER, SvO2 e tendencia."],
            ["Lactato final > 4 mmol/L", "Alto", "Procurar acidose metabolica e causas globais/regionais."],
            ["Razao lactato final/inicial > 1,1", "Moderado", "Pode alertar mesmo se o valor final ainda for baixo."],
            ["Gap PCO2 >= 6 mmHg", "Moderado", "Pode sugerir componente de fluxo/perfusao regional, sempre contextualizar."],
            ["pH/HCO3/PaCO2", "Resumo acido-base", "Classificar o disturbio final e checar compensacoes."],
        ],
        widths=[2.1, 1.8, 3.7],
    )

    add_heading(doc, "Hemodiluicao, Hb e PBM", 1)
    doc.add_paragraph(
        "A hemoglobina e um dos maiores determinantes do CaO2. Quando Hb cai, o iDO2 pode cair mesmo que o fluxo permaneca igual. Ainda assim, transfusao nunca deve ser indicada por um unico numero isolado."
    )
    add_bullets(
        doc,
        [
            "O PerfuseLab alerta Hb < 7,5 g/dL como atencao PBM.",
            "O alerta nao e ordem de transfusao. Ele pede correlacao com iDO2, lactato, SvO2, sangramento, reserva do paciente e protocolo.",
            "Hemodiluicao tambem pode alterar sodio, cloreto, calcio, pressao oncótica e equilibrio acido-base.",
            "Casos 02, 08 e 18 sao bons para treinar queda de Hb/Hct e impacto no iDO2.",
        ],
    )

    add_heading(doc, "Temperatura e Alvo Termico Local", 1)
    doc.add_paragraph(
        "O PerfuseLab exibe o alvo GDP fixo de 280 mL/min/m2 e tambem uma tabela termica local configuravel. A hipotermia reduz consumo, mas isso nao deve virar leitura automatica sem contexto."
    )
    add_table(
        doc,
        ["Temperatura", "Alvo termico local exibido", "Observacao"],
        [
            ["25 C", "110 mL/min/m2", "Valor de protocolo local."],
            ["28 C", "152 mL/min/m2", "Interpolado entre pontos."],
            ["30 C", "180 mL/min/m2", "Interpretar com alfa-stat/pH-stat."],
            ["33 C", "223 mL/min/m2", "Reaquecimento aumenta demanda."],
            ["35 C", "252 mL/min/m2", "Checar tendencia de lactato e O2ER."],
            ["37 C", "280 mL/min/m2", "Converge para alvo GDP fixo."],
        ],
        widths=[1.3, 2.0, 4.3],
    )

    add_heading(doc, "Alertas e Risco Integrado do PerfuseLab", 1)
    doc.add_paragraph(
        "Os alertas sao baseados em regras locais do codigo. O risco integrado resume a quantidade e gravidade desses alertas; ele nao e EuroSCORE, STS, nem modelo prognostico validado."
    )
    add_table(
        doc,
        ["Categoria", "Regra local", "Peso no risco"],
        [
            ["iDO2", "< 90% de 280 = alto; 90 a <100% = moderado", "Alto ou moderado"],
            ["O2ER", "> 40% alto; >= 30% moderado", "Alto ou moderado"],
            ["DO2/VO2", "< 2,5 alto; <= 3,33 moderado", "Alto ou moderado"],
            ["SvO2", "< 65% alto; <= 70% moderado", "Alto ou moderado"],
            ["PAM", "< 60 mmHg", "Moderado"],
            ["Gap PCO2", ">= 6 mmHg", "Moderado"],
            ["K", "< 3,5 mmol/L", "Moderado"],
            ["Glicose", "> 180 mg/dL", "Moderado"],
            ["Hb", "< 7,5 g/dL", "Moderado"],
            ["Lactato", "Final > 4 alto; final >= 2 moderado; razao > 1,1 moderado", "Alto ou moderado"],
            ["Tempo de CEC", "60 a 120 min moderado; > 120 min alto", "Alto ou moderado"],
            ["TCA", "< 400 s alto; < 480 s moderado, quando informado", "Alto ou moderado"],
            ["Informativos", "PaO2 alta, iDO2 a conferir, marcadores discordantes", "Nao elevam risco"],
        ],
        widths=[1.5, 4.4, 1.7],
    )
    add_callout(
        doc,
        "Como o risco final e decidido",
        "Se houver qualquer alerta alto, o risco integrado vira ALTO. Se nao houver alerta alto, mas existirem dois ou mais alertas moderados, vira MODERADO. Com menos de dois moderados e nenhum alto, fica BAIXO.",
        LIGHT_BLUE,
    )

    add_heading(doc, "Como Interpretar a Tabela de Monitorizacao", 1)
    add_numbered(
        doc,
        [
            "Leia a linha inicial: ela e seu ponto de comparacao.",
            "Procure o pior ponto de iDO2, nao apenas o valor final.",
            "Veja se o deficit foi profundo e curto ou raso e prolongado.",
            "Compare iDO2 com O2ER e SvO2: oferta baixa com extracao alta e mais preocupante.",
            "Cheque lactato, gap PCO2 e acido-base: eles dizem se a fisiologia esta cobrando a conta.",
            "Volte para Hb/Hct, fluxo/IC, SaO2 e PaO2 para entender o mecanismo.",
            "Por fim, leia PAM, temperatura, potassio, sodio, glicose e observacoes operacionais.",
        ],
    )

    add_heading(doc, "Como Interpretar os Graficos", 1)
    add_table(
        doc,
        ["Grafico", "Padrao tranquilizador", "Padrao que merece estudo"],
        [
            ["iDO2", "Acima de 280 e estavel.", "Queda progressiva, vale abaixo de 280, ou recuperacao tardia."],
            ["Lactato", "Estavel ou queda.", "Subida progressiva, especialmente > 2 ou > 4 mmol/L."],
            ["Hb/Hct", "Queda pequena e coerente com prime/volume.", "Queda brusca com iDO2 limitrofe."],
            ["SvO2/O2ER", "SvO2 preservada e O2ER baixa.", "SvO2 baixa, O2ER >= 30%, DO2/VO2 <= 3,33."],
            ["Gap PCO2", "Baixo e estavel.", "Aumento progressivo, principalmente com lactato subindo."],
        ],
        widths=[1.5, 3.0, 3.1],
    )

    add_heading(doc, "Exemplo Guiado: Caso 10", 1)
    doc.add_paragraph(
        "O caso 10 foi construido para ser matematicamente coerente com o caso valvar original. Ele tem quatro tempos alinhados, SaO2 medida e iDO2 recalculavel."
    )
    add_table(
        doc,
        ["Tempo", "Hb", "Fluxo/IC", "iDO2", "Lactato", "Leitura"],
        [
            ["0 min", "12,0", "IC 2,4", "406,77", "1,44", "Oferta muito preservada no inicio."],
            ["35 min", "8,2", "IC 2,5", "294,01", "1,55", "Hemodiluicao importante, mas iDO2 ainda acima de 280."],
            ["60 min", "7,8", "IC 2,5", "281,91", "2,11", "Ponto mais proximo do alvo GDP."],
            ["90 min", "8,7", "IC 2,7", "328,34", "2,55", "Recupera oferta, mas lactato e eletrólitos merecem atencao."],
        ],
        widths=[1.0, 0.8, 1.2, 1.1, 1.0, 2.5],
    )
    add_callout(
        doc,
        "Pergunta central do caso 10",
        "A queda de Hb reduziu o conteudo arterial, mas o aumento do indice de fluxo manteve o iDO2 proximo/acima do alvo. Ainda assim, lactato, potassio e glicose finais impedem uma leitura simplista.",
        LIGHT_GREEN,
    )

    add_heading(doc, "Trilha dos 20 Casos Clinicos", 1)
    cases = [
        ["01", "CEC estavel", "Moderado", "Reconhecer padrao globalmente tranquilo e efeito do tempo de CEC."],
        ["02", "Hemodiluicao com baixo iDO2", "Alto", "Ver como Hb baixa derruba CaO2 e iDO2."],
        ["03", "Recuperacao com ajuste de fluxo", "Alto", "Aprender que recuperacao final nao apaga deficit acumulado."],
        ["04", "Lactato crescente com iDO2 preservado", "Alto", "Separar metabolismo de oferta global."],
        ["05", "Extracao elevada e SvO2 baixa", "Alto", "Identificar consumo/extracao em zona critica."],
        ["06", "Acidose respiratoria", "Moderado", "Diferenciar gasometria ventilatoria de baixo iDO2."],
        ["07", "Hiperoxia com oferta adequada", "Moderado", "Entender PaO2 alta como informativo, nao prescricao."],
        ["08", "CEC prolongada, anemia e lactato", "Alto", "Integrar tempo, hemodiluicao e lactato."],
        ["09", "Hipotermia e reaquecimento", "Alto", "Comparar alvo fixo e alvo termico local."],
        ["10", "Reoperacao valvar coerente", "Moderado", "Auditar formulas e alinhamento temporal."],
        ["11", "PAM baixa com iDO2 preservado", "Moderado", "Nao confundir oferta global com pressao adequada."],
        ["12", "Hiperglicemia/lactato com iDO2 preservado", "Moderado", "Reconhecer estresse metabolico sem baixo iDO2."],
        ["13", "Hipocalemia no reaquecimento", "Moderado", "Ler eletrólitos alem dos graficos principais."],
        ["14", "Obesidade: fluxo total alto, IC baixo", "Alto", "Indexar por BSA antes de confiar no fluxo."],
        ["15", "iDO2 informado sem SaO2", "Alto", "Distinguir valor informado de valor recalculado."],
        ["16", "Gap PCO2 alto com SvO2 preservada", "Moderado", "Estudar marcadores discordantes."],
        ["17", "Hipercalemia e limite do framework", "Moderado", "Perceber achados relevantes ainda nao alertados."],
        ["18", "Hiponatremia e hemodiluicao", "Moderado", "Entender Hb, sodio e iDO2 limitrofe."],
        ["19", "Acidose metabolica com lactato critico", "Alto", "Valorizar lactato > 4 mesmo com iDO2 bom."],
        ["20", "Queda de SaO2 arterial recuperada", "Moderado", "Ver efeito da saturacao arterial no CaO2."],
    ]
    add_table(doc, ["Caso", "Tema", "Risco", "Objetivo de estudo"], cases, widths=[0.7, 2.3, 1.0, 3.6])

    add_heading(doc, "Armadilhas Comuns", 1)
    add_bullets(
        doc,
        [
            "Confundir fluxo total com fluxo indexado.",
            "Usar lactato em mg/dL como se fosse mmol/L.",
            "Interpretar PaO2 alta como grande aumento de CaO2; na formula, Hb e SaO2 pesam muito mais.",
            "Ignorar o tempo abaixo do alvo quando o valor final ja melhorou.",
            "Achar que risco BAIXO significa caso perfeito. Pode haver dado nao contemplado pela regra.",
            "Achar que risco ALTO significa uma conduta automatica. Ele significa necessidade de leitura contextual.",
            "Confiar em iDO2 informado sem conferir se SaO2, Hb e IC permitiriam recalculo.",
        ],
    )

    add_heading(doc, "Checklist de Interpretacao", 1)
    add_table(
        doc,
        ["Item", "Pergunta rapida", "Resposta no caso"],
        [
            ["Unidade", "Lactato esta em mmol/L ou mg/dL?", ""],
            ["Auditoria", "Tenho SaO2 medida para recalcular iDO2?", ""],
            ["Tamanho", "Qual BSA e IC?", ""],
            ["Oferta", "Qual pior iDO2 e por quanto tempo ficou abaixo de 280?", ""],
            ["Extracao", "O2ER >= 30%, DO2/VO2 <= 3,33 ou SvO2 <= 70%?", ""],
            ["Metabolismo", "Lactato subiu? pH/HCO3 sugerem acidose?", ""],
            ["Regional", "Gap PCO2 esta >= 6 ou subindo?", ""],
            ["Operacional", "PAM, temperatura, glicose, K, Na e Hb fazem sentido?", ""],
            ["Conclusao", "O risco integrado combina com minha leitura?", ""],
        ],
        widths=[1.2, 4.1, 2.3],
    )

    add_heading(doc, "Glossario Curto", 1)
    add_table(
        doc,
        ["Termo", "Significado"],
        [
            ["BSA/SC", "Superficie corporal. Usada para indexar fluxo e oferta."],
            ["IC", "Indice cardiaco ou indice de fluxo em CEC, L/min/m2."],
            ["CaO2", "Conteudo arterial de oxigenio."],
            ["iDO2", "Oferta de oxigenio indexada."],
            ["iVO2", "Consumo de oxigenio indexado, quando informado."],
            ["O2ER", "Taxa de extracao de oxigenio."],
            ["SvO2", "Saturacao venosa mista/venosa do circuito, conforme coleta."],
            ["Gap PCO2", "Diferenca venoarterial de CO2. Ajuda a estudar componente de fluxo/perfusao."],
            ["AUC", "Area acumulada de deficit abaixo de um alvo."],
            ["GDP", "Goal-directed perfusion, estrategia orientada por metas como iDO2."],
        ],
        widths=[1.4, 6.2],
    )

    add_heading(doc, "Referencias de Apoio", 1)
    add_reference(
        doc,
        "Mosteller RD. Simplified calculation of body-surface area. N Engl J Med. 1987.",
        "https://pubmed.ncbi.nlm.nih.gov/3657876/",
    )
    add_reference(
        doc,
        "Ranucci et al. Goal-directed perfusion trial, alvo de iDO2 >= 280 mL/min/m2.",
        "https://pubmed.ncbi.nlm.nih.gov/29778331/",
    )
    add_reference(
        doc,
        "EACTS/EACTAIC/EBCP Guidelines on cardiopulmonary bypass in adult cardiac surgery.",
        "https://pmc.ncbi.nlm.nih.gov/articles/PMC11826095/",
    )
    add_reference(
        doc,
        "EACTS/EACTAIC patient blood management guideline in adult cardiac surgery.",
        "https://pmc.ncbi.nlm.nih.gov/articles/PMC12256976/",
    )
    add_reference(
        doc,
        "Formula operacional contemporanea de iDO2 em CEC e discussao de perfusao dirigida por oxigenio.",
        "https://pmc.ncbi.nlm.nih.gov/articles/PMC12039435/",
    )
    add_reference(
        doc,
        "Formulas de CaO2, CvO2, DO2 e VO2 em monitorizacao hemodinamica.",
        "https://pmc.ncbi.nlm.nih.gov/articles/PMC10488867/",
    )
    add_reference(
        doc,
        "STS/SCA/AmSECT guideline de anticoagulacao durante CEC.",
        "https://pmc.ncbi.nlm.nih.gov/articles/PMC5850589/",
    )

    doc.add_paragraph()
    add_callout(
        doc,
        "Ultima mensagem para estudo",
        "A boa interpretacao nao e decorar o alerta. E contar a historia: o que caiu, o que compensou, por quanto tempo, com que custo metabolico, e quais dados ainda faltam.",
        LIGHT_GREEN,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_manual()
